"""
Report generator: assembles Markdown and HTML from analysis artifacts.
"""
from typing import Dict, Any, List
from datetime import datetime
import json
import io


def generate_markdown(payload: Dict[str, Any]) -> str:
    ds = payload.get("dataset", {})
    ethics = payload.get("ethics", {})
    eda = payload.get("eda", {})
    story = payload.get("story", {})
    ml = payload.get("ml", {})

    md = []
    md.append(f"# AETHER Insight Report\n")
    md.append(f"_Generated: {datetime.utcnow().isoformat()}Z_\n")
    md.append("\n## Dataset Overview\n")
    md.append(f"- Name: {ds.get('name','N/A')}\n- Rows: {ds.get('row_count','?')}\n- Columns: {ds.get('column_count','?')}\n")

    md.append("\n## Ethical Analysis\n")
    if ethics:
        md.append("- Sensitive Attributes: " + ", ".join(ethics.get("sensitive_attributes", [])) + "\n")
        for n in ethics.get("notes", []):
            md.append(f"- {n}\n")
    else:
        md.append("No ethical notes available.\n")

    md.append("\n## EDA Highlights\n")
    for ins in eda.get("insights", []):
        md.append(f"- {ins}\n")

    md.append("\n## Story\n")
    for s in story.get("sections", []):
        md.append(f"### {s.get('title')}\n{s.get('body')}\n\n")

    if ml:
        md.append("\n## Modeling Results\n")
        md.append(f"Best Model: {ml.get('best_model','N/A')}\n\n")
        best = ml.get("best_score", {})
        for k, v in best.items():
            if isinstance(v, (int, float)):
                md.append(f"- {k}: {v}\n")

    md.append("\n## Recommendations\n")
    md.append("- Address data quality issues before deployment.\n")
    md.append("- Monitor fairness metrics periodically.\n")
    return "".join(md)


def markdown_to_html(md: str) -> str:
    # Minimal conversion for portability (avoid heavy deps)
    import html
    lines = md.split("\n")
    html_lines: List[str] = []
    for line in lines:
        if line.startswith("# "):
            html_lines.append(f"<h1>{html.escape(line[2:])}</h1>")
        elif line.startswith("## "):
            html_lines.append(f"<h2>{html.escape(line[3:])}</h2>")
        elif line.startswith("### "):
            html_lines.append(f"<h3>{html.escape(line[4:])}</h3>")
        elif line.startswith("- "):
            html_lines.append(f"<li>{html.escape(line[2:])}</li>")
        else:
            html_lines.append(f"<p>{html.escape(line)}</p>")
    return "\n".join(html_lines)


def _render_preview_table(rows: List[Dict[str, Any]]) -> str:
    if not rows:
        return ""
    cols = list(rows[0].keys())
    # head rows only
    head = rows[:10]
    header_html = "".join([f"<th style='padding:6px 10px;border-bottom:1px solid #1e293b;text-align:left'>{c}</th>" for c in cols])
    body_html = []
    for r in head:
        tds = "".join([f"<td style='padding:6px 10px;border-bottom:1px solid #0b1220'>{str(r.get(c,'')).replace('<','&lt;').replace('>','&gt;')}</td>" for c in cols])
        body_html.append(f"<tr>{tds}</tr>")
    return f"<table style='width:100%;border-collapse:collapse;margin:8px 0 16px 0'><thead><tr>{header_html}</tr></thead><tbody>{''.join(body_html)}</tbody></table>"


def _collect_plot_specs(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    eda = payload.get("eda", {})
    ml = payload.get("ml", {})
    plots: List[Dict[str, Any]] = []
    for trace in eda.get("histograms", [])[:4]:
        plots.append({"title": trace.get("title", "Histogram"), "fig": trace})
    corr = eda.get("correlations")
    if corr:
        plots.append({"title": "Correlation Heatmap", "fig": corr})
    for box in eda.get("box_plots", [])[:4]:
        plots.append({"title": box.get("title", "Boxplot"), "fig": box})
    if ml:
        best = ml.get("best_model")
        if best:
            metrics = ml.get("best_score", {})
            roc = metrics.get("roc_curve")
            if roc:
                plots.append({
                    "title": f"ROC Curve — {best}",
                    "fig": {
                        "data": [{"type": "scatter", "mode": "lines", "x": roc.get("fpr", []), "y": roc.get("tpr", []), "name": "ROC"}, {"type": "scatter", "mode": "lines", "x": [0,1], "y": [0,1], "name": "Random", "line": {"dash": "dash"}}],
                        "layout": {"xaxis": {"title": "FPR"}, "yaxis": {"title": "TPR"}}
                    }
                })
            cm = metrics.get("confusion_matrix")
            if cm:
                plots.append({
                    "title": f"Confusion Matrix — {best}",
                    "fig": {
                        "data": [{"type": "heatmap", "z": cm}],
                        "layout": {"xaxis": {"title": "Predicted"}, "yaxis": {"title": "Actual"}}
                    }
                })
    return plots


def generate_html(payload: Dict[str, Any]) -> str:
    """Generate a standalone HTML report with embedded Plotly visuals."""
    md_html = markdown_to_html(generate_markdown(payload))
    ds = payload.get("dataset", {})

    plots = _collect_plot_specs(payload)

    # Data preview table if provided
    preview_rows = ds.get("preview") or []
    preview_html = _render_preview_table(preview_rows)

    # Build HTML
    plot_divs = []
    for idx, p in enumerate(plots):
        div_id = f"plot_{idx}"
        fig_json = json.dumps(p["fig"])  # fig is either a full fig or a {data, layout}
        script = f"""
        <div class=\"section\"><h3>{p['title']}</h3><div id=\"{div_id}\" style=\"height:420px;\"></div></div>
        <script>
          (function() {{
            var fig = {fig_json};
            var data = fig.data || [fig];
            var layout = fig.layout || {{ title: {json.dumps(p['title'])} }};
            window.Plotly.newPlot('{div_id}', data, layout, {{displayModeBar: false}});
          }})();
        </script>
        """
        plot_divs.append(script)

    html = f"""
    <!doctype html>
    <html>
    <head>
      <meta charset=\"utf-8\" />
      <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />
      <title>AETHER Insight Report</title>
      <script src=\"https://cdn.plot.ly/plotly-2.26.0.min.js\"></script>
      <style>
        body {{ font-family: Inter, ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, 'Apple Color Emoji', 'Segoe UI Emoji'; padding: 24px; background: #0b1220; color: #e5eefc; }}
        h1, h2, h3 {{ color: #eaf2ff; }}
        .section {{ background: #0f172a; border: 1px solid #1e293b; border-radius: 12px; padding: 16px 20px; margin: 16px 0; }}
        table {{ font-size: 12px; }}
      </style>
    </head>
    <body>
      <div class=\"section\">{md_html}</div>
      {('<h2>Data Preview</h2>' + preview_html) if preview_html else ''}
      <h2>Visuals</h2>
      {''.join(plot_divs)}
    </body>
    </html>
    """
    return html


def _export_plot_images(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Render plot specs to PNG bytes using plotly + kaleido."""
    import plotly.graph_objects as go
    import plotly.io as pio
    specs = _collect_plot_specs(payload)
    images: List[Dict[str, Any]] = []
    for s in specs:
        fig_def = s["fig"]
        if "data" in fig_def:
            fig = go.Figure(fig_def.get("data", []), fig_def.get("layout", {}))
        else:
            # single-trace shorthand
            fig = go.Figure([fig_def], {})
        img_bytes = pio.to_image(fig, format='png', width=900, height=500, scale=2)
        images.append({"title": s["title"], "image_data": img_bytes, "png": img_bytes})
    return images


def generate_docx(payload: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches
    import io
    doc = Document()
    doc.add_heading('AETHER Insight Report', level=1)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")

    ds = payload.get("dataset", {})
    doc.add_heading('Dataset Overview', level=2)
    doc.add_paragraph(f"Name: {ds.get('name','N/A')}")
    doc.add_paragraph(f"Rows: {ds.get('row_count','?')} | Columns: {ds.get('column_count','?')}")

    # Data preview table
    preview = ds.get("preview") or []
    if preview:
        cols = list(preview[0].keys())
        table = doc.add_table(rows=1, cols=len(cols))
        hdr = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr[i].text = str(c)
        for r in preview[:10]:
            row_cells = table.add_row().cells
            for i, c in enumerate(cols):
                row_cells[i].text = str(r.get(c, ''))

    # Ethics
    ethics = payload.get("ethics", {})
    doc.add_heading('Ethical Analysis', level=2)
    if ethics:
        doc.add_paragraph("Sensitive Attributes: " + ", ".join(ethics.get("sensitive_attributes", [])))
        for n in ethics.get("notes", []):
            doc.add_paragraph("- " + n)

    # EDA insights
    eda = payload.get("eda", {})
    if eda.get("insights"):
        doc.add_heading('EDA Highlights', level=2)
        for ins in eda.get("insights", []):
            doc.add_paragraph("- " + str(ins))

    # Story
    story = payload.get("story", {})
    if story.get("sections"):
        doc.add_heading('Story', level=2)
        for s in story.get("sections", []):
            if s.get('title'):
                doc.add_heading(s.get('title'), level=3)
            if s.get('body'):
                doc.add_paragraph(s.get('body'))

    # ML results
    ml = payload.get("ml", {})
    if ml:
        doc.add_heading('Modeling Results', level=2)
        doc.add_paragraph(f"Best Model: {ml.get('best_model','N/A')}")
        best = ml.get("best_score", {})
        for k, v in best.items():
            if isinstance(v, (int, float)):
                doc.add_paragraph(f"- {k}: {v}")

    # Visuals as images
    images = _export_plot_images(payload)
    if images:
        doc.add_heading('Visuals', level=2)
        for im in images:
            doc.add_paragraph(im["title"])
            run = doc.add_paragraph().add_run()
            run.add_picture(io.BytesIO(im["png"]), width=Inches(6.5))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_pdf(payload: Dict[str, Any]) -> bytes:
    """
    Generate professional PDF report with colorful EDA visuals
    Professional format with proper sections, colors, and embedded visualizations
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.colors import HexColor
    import io
    import base64
    import json
    
    width, height = A4
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    
    # Color scheme
    primary_color = HexColor('#2563eb')
    secondary_color = HexColor('#10b981')
    accent_color = HexColor('#f59e0b')
    text_color = HexColor('#0f172a')
    light_gray = HexColor('#f8fafc')
    
    y = height - 40
    
    def draw_header(title: str, subtitle: str = ""):
        nonlocal y
        if y < height - 200:
            c.showPage()
            y = height - 40
        
        # Header background
        c.setFillColor(primary_color)
        c.rect(0, y - 30, width, 50, fill=1, stroke=0)
        
        # Title
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 20)
        c.drawString(50, y, title)
        
        if subtitle:
            c.setFont("Helvetica", 10)
            c.drawString(50, y - 18, subtitle)
        
        y -= 60
    
    def write_section(title: str, color=primary_color):
        nonlocal y
        if y < 150:
            c.showPage()
            y = height - 40
        
        # Section header with colored background
        c.setFillColor(color)
        c.rect(40, y - 20, width - 80, 25, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y - 15, title)
        y -= 35
    
    def write_text(text: str, size=10, bold=False, indent=0):
        nonlocal y
        if y < 80:
            c.showPage()
            y = height - 40
        
        c.setFillColor(text_color)
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        # Handle text wrapping
        words = text.split()
        line = ""
        for word in words:
            test_line = line + word + " "
            if c.stringWidth(test_line, "Helvetica-Bold" if bold else "Helvetica", size) > width - 100 - indent:
                if line:
                    c.drawString(50 + indent, y, line)
                    y -= size + 4
                line = word + " "
            else:
                line = test_line
        if line:
            c.drawString(50 + indent, y, line)
            y -= size + 4
    
    def add_image_from_base64(img_data: bytes, width_inch=6):
        nonlocal y
        if y < 200:
            c.showPage()
            y = height - 40
        
        try:
            img = ImageReader(io.BytesIO(img_data))
            img_width, img_height = img.getSize()
            aspect = img_height / img_width
            display_width = width_inch * inch
            display_height = display_width * aspect
            
            if y - display_height < 50:
                c.showPage()
                y = height - 40
            
            c.drawImage(img, 50, y - display_height, width=display_width, height=display_height)
            y -= display_height + 20
        except Exception as e:
            write_text(f"[Image could not be loaded: {str(e)}]", size=8)
    
    # Cover page
    draw_header("AETHER Insight Platform", "Data Analysis Report")
    y -= 20
    
    ds = payload.get("dataset", {})
    write_text(f"Dataset: {ds.get('name', 'N/A')}", size=14, bold=True)
    write_text(f"Rows: {ds.get('row_count', '?')} | Columns: {ds.get('column_count', '?')}", size=12)
    write_text(f"Generated: {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}", size=10)
    
    # Dataset Overview
    c.showPage()
    y = height - 40
    write_section("1. Dataset Overview", primary_color)
    write_text(f"Dataset Name: {ds.get('name', 'N/A')}", bold=True)
    write_text(f"Total Rows: {ds.get('row_count', '?')}")
    write_text(f"Total Columns: {ds.get('column_count', '?')}")
    
    # Ethical Analysis
    write_section("2. Ethical Analysis & GDPR Compliance", accent_color)
    ethics = payload.get("ethics", {})
    if ethics:
        if ethics.get("sensitive_attributes"):
            write_text("Sensitive Attributes Detected:", bold=True)
            write_text(", ".join(ethics.get("sensitive_attributes", [])), indent=20)
        
        if ethics.get("pii_detected"):
            write_text("PII Detected:", bold=True)
            for pii_type, cols in ethics.get("pii_detected", {}).items():
                write_text(f"  • {pii_type}: {', '.join(cols)}", indent=20)
        
        if ethics.get("gdpr_compliance"):
            gdpr = ethics.get("gdpr_compliance", {})
            score = gdpr.get("compliance_score", 100)
            write_text(f"GDPR Compliance Score: {score}/100", bold=True)
            if gdpr.get("recommendations"):
                write_text("Recommendations:", bold=True)
                for rec in gdpr.get("recommendations", [])[:5]:
                    write_text(f"  • {rec}", indent=20)
        
        if ethics.get("notes"):
            write_text("Key Notes:", bold=True)
            for note in ethics.get("notes", [])[:6]:
                write_text(f"  • {note}", indent=20)
    
    # EDA Section with Visuals
    eda = payload.get("eda", {})
    if eda:
        write_section("3. Exploratory Data Analysis", secondary_color)
        
        # Data type
        if eda.get("data_type"):
            write_text(f"Detected Data Type: {eda.get('data_type', 'mixed').replace('_', ' ').title()}", bold=True)
            y -= 10
        
        # Insights
        if eda.get("insights"):
            write_text("Key Insights:", bold=True)
            for ins in eda.get("insights", [])[:8]:
                write_text(f"  • {ins}", indent=20)
        
        # Visualizations
        images = _export_plot_images(payload)
        if images:
            write_text("Visualizations:", bold=True)
            y -= 10
            for img_info in images[:10]:  # Limit to 10 images
                if img_info.get("image_data"):
                    add_image_from_base64(img_info["image_data"], width_inch=6)
    
    # Story Section
    story = payload.get("story", {})
    if story and story.get("sections"):
        write_section("4. Data Story", primary_color)
        for section in story.get("sections", []):
            if section.get("title"):
                write_text(section.get("title"), size=12, bold=True)
            if section.get("body"):
                write_text(section.get("body"), size=10)
            y -= 10
    
    # ML Results (if available)
    ml = payload.get("ml", {})
    if ml:
        write_section("5. Machine Learning Analysis", accent_color)
        write_text(f"Best Model: {ml.get('best_model', 'N/A')}", bold=True)
        best = ml.get("best_score", {})
        if best:
            write_text("Performance Metrics:", bold=True)
            for k, v in best.items():
                if isinstance(v, (int, float)):
                    write_text(f"  • {k}: {v:.4f}" if isinstance(v, float) else f"  • {k}: {v}", indent=20)
    
    # Footer
    c.showPage()
    y = height - 40
    c.setFillColor(light_gray)
    c.rect(0, 0, width, 40, fill=1, stroke=0)
    c.setFillColor(text_color)
    c.setFont("Helvetica", 8)
    c.drawString(50, 15, f"AETHER Insight Platform - Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    
    c.save()
    return bio.getvalue()
